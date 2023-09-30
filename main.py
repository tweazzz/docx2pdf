import pandas as pd
import csv
from docx import Document
from docx.api import Document
from docx2pdf import convert
from pathlib import Path

# pip install pandas python-docx docx2pdf openpyxl


# Converting xlsx to csv ////
read_file = pd.read_excel("Сотрудники.xlsx")
read_file.to_csv ("Сотрудники.csv", encoding='utf-8-sig',index = None,header=True)


with open("Сотрудники.csv",encoding='utf-8-sig') as file:
    key = [line.strip() for line in file.readline().split(',')]
    value = [line.strip().split(',') for line in file.readlines()]
    res = [dict(zip(key,i)) for i in value]

full_name = []
for i in value:
    full_name.append(i[0]+' '+i[1]+' '+i[2])


def conver2pdf():
    document = Document("Шаблон.docx")
    table = document.tables[0]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{person_FIO}}' in paragraph.text:
                        for j in range(len(full_name)):
                            paragraph.text = str(full_name[j])
                            my_file = Path(f"Доска почета/docx/{full_name[j]}.docx")
                            my_pdf = Path(f'Доска почета/{full_name[j]}.pdf')
                            if not my_file.is_file():
                                document.save(f'Доска почета/docx/{full_name[j]}.docx')
                                convert(f'Доска почета/docx/{full_name[j]}.docx', f'Доска почета/{full_name[j]}.pdf')
                                print(f'[INFO]: File {full_name[j]} is added!')
                            else:
                                if not my_pdf.is_file():
                                    convert(f'Доска почета/docx/{full_name[j]}.docx', f'Доска почета/{full_name[j]}.pdf')
                                    print(f'[INFO]: File {full_name[j]} converted to pdf format!')
                                    break
                                print(f'[INFO]: File {full_name[j]} is already exists!!!')

def main():
    conver2pdf()


if __name__ == "__main__":
    main()